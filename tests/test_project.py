import os
from datetime import datetime
from pathlib import Path

import docx
import pytest

from videoai.core.project import (
    brief_prompt,
    load_brief,
    read_brief,
    resolve_clip_dir,
    snapshot_output,
)


def test_resolve_clip_dir_prefers_input(tmp_path: Path):
    (tmp_path / "input").mkdir()
    (tmp_path / "video").mkdir()
    assert resolve_clip_dir(tmp_path) == tmp_path / "input"


def test_resolve_clip_dir_falls_back_to_video(tmp_path: Path):
    (tmp_path / "video").mkdir()
    assert resolve_clip_dir(tmp_path) == tmp_path / "video"


def test_resolve_clip_dir_falls_back_to_project_dir(tmp_path: Path):
    assert resolve_clip_dir(tmp_path) == tmp_path


def test_read_brief_returns_empty_string_when_nothing_present(tmp_path: Path):
    assert read_brief(tmp_path) == ""


def test_read_brief_includes_project_yaml_and_notes(tmp_path: Path):
    (tmp_path / "project.yaml").write_text("title: Slime review\n", encoding="utf-8")
    (tmp_path / "notes.md").write_text("Keep the laugh at the end.", encoding="utf-8")

    brief = read_brief(tmp_path)

    assert "Slime review" in brief
    assert "Keep the laugh at the end." in brief


def test_read_brief_reads_docx_from_description(tmp_path: Path):
    description = tmp_path / "description"
    description.mkdir()
    document = docx.Document()
    document.add_paragraph("Pimple Popping Stress Toy")
    document.add_paragraph("Refillable, two in one.")
    document.save(description / "product.docx")

    brief = read_brief(tmp_path)

    assert "Pimple Popping Stress Toy" in brief
    assert "Refillable, two in one." in brief


def test_read_brief_reads_markdown_and_text_from_description(tmp_path: Path):
    description = tmp_path / "description"
    description.mkdir()
    (description / "a.md").write_text("markdown content", encoding="utf-8")
    (description / "b.txt").write_text("plain content", encoding="utf-8")

    brief = read_brief(tmp_path)

    assert "markdown content" in brief
    assert "plain content" in brief


def test_read_brief_ignores_macos_metadata_and_images(tmp_path: Path):
    description = tmp_path / "description"
    description.mkdir()
    (description / ".DS_Store").write_bytes(b"junk")
    (description / "._notes.md").write_bytes(b"junk")
    (description / "photo.jpeg").write_bytes(b"\xff\xd8\xff")
    (description / "real.md").write_text("real content", encoding="utf-8")

    brief = read_brief(tmp_path)

    assert brief.strip() == "real content"


def test_an_unreadable_brief_file_stops_the_run_naming_it(tmp_path: Path):
    """A brief that cannot be read is not a shorter brief. It used to become the
    literal text `[could not read broken.docx]` in the prompt, so the arc and the
    must-include list vanished and every stage planned as if nothing was asked."""
    description = tmp_path / "description"
    description.mkdir()
    (description / "broken.docx").write_bytes(b"not a real docx")
    (description / "real.md").write_text("real content", encoding="utf-8")

    with pytest.raises(RuntimeError) as error:
        read_brief(tmp_path)

    assert "broken.docx" in str(error.value)


# --- The brief as fields: `arc`, `must_include` and the rest used to be
# conventions honoured only by a model's goodwill ---


def test_a_filled_in_brief_parses_into_fields(tmp_path: Path):
    (tmp_path / "project.yaml").write_text(
        "title: Slime review\n"
        "toy_name: Refillable Slime Blob\n"
        "toy_description: A squishy blob you fill with paint.\n"
        "summary: Ten minutes of popping, cut to three.\n"
        "target_duration_seconds: 180\n"
        "arc:\n  - opens the box\n  - pops the pimples\n"
        "must_include:\n  - the first pop\n"
        "avoid:\n  - the delivery label\n"
        "style: playful\n",
        encoding="utf-8",
    )

    brief = load_brief(tmp_path)

    assert brief.title == "Slime review"
    assert brief.subject == "Refillable Slime Blob"
    assert brief.subject_description == "A squishy blob you fill with paint."
    assert brief.summary == "Ten minutes of popping, cut to three."
    assert brief.arc == ["opens the box", "pops the pimples"]
    assert brief.must_include == ["the first pop"]
    assert brief.avoid == ["the delivery label"]
    assert brief.target_duration_seconds == 180.0


def test_the_documented_example_brief_parses(tmp_path: Path):
    """The file the README tells creators to copy has to be the one that parses."""
    import shutil

    example = Path(__file__).resolve().parents[1] / "docs" / "project-yaml-example.yaml"
    shutil.copy(example, tmp_path / "project.yaml")

    brief = load_brief(tmp_path)

    assert brief.subject.startswith("2-in-1 Refillable")
    assert len(brief.arc) == 4
    assert len(brief.must_include) == 3
    assert brief.target_duration_seconds == 180.0


def test_free_prose_in_the_brief_still_reaches_the_model(tmp_path: Path):
    """Nothing parses notes.md or description/, and nothing should — but a brief
    that is only prose must not become an empty brief."""
    (tmp_path / "notes.md").write_text("Keep the laugh at the end.", encoding="utf-8")

    brief = load_brief(tmp_path)

    assert brief.arc == []
    assert "Keep the laugh at the end." in brief.raw
    assert "Keep the laugh at the end." in brief_prompt(brief)


def test_an_unknown_key_survives_in_the_raw_brief(tmp_path: Path):
    (tmp_path / "project.yaml").write_text(
        "title: T\nsomething_new: the parent must never be the presenter\n",
        encoding="utf-8",
    )

    assert "the parent must never be the presenter" in brief_prompt(load_brief(tmp_path))


def test_a_brief_that_is_not_a_mapping_is_not_an_error(tmp_path: Path):
    (tmp_path / "project.yaml").write_text("just some prose about a toy\n", encoding="utf-8")

    brief = load_brief(tmp_path)

    assert brief.arc == []
    assert "just some prose about a toy" in brief.raw


def test_unparseable_yaml_stops_the_run_rather_than_being_ignored(tmp_path: Path):
    (tmp_path / "project.yaml").write_text("arc:\n  - one\n - two\n", encoding="utf-8")

    with pytest.raises(RuntimeError) as error:
        load_brief(tmp_path)

    assert "project.yaml" in str(error.value)


def test_a_target_duration_that_is_not_a_number_is_refused(tmp_path: Path):
    (tmp_path / "project.yaml").write_text(
        "target_duration_seconds: three minutes\n", encoding="utf-8"
    )

    with pytest.raises(RuntimeError, match="three minutes"):
        load_brief(tmp_path)


def test_an_arc_written_as_lines_rather_than_a_yaml_list_still_parses(tmp_path: Path):
    (tmp_path / "project.yaml").write_text(
        "arc: |\n  - opens the box\n  - pops the pimples\n", encoding="utf-8"
    )

    assert load_brief(tmp_path).arc == ["opens the box", "pops the pimples"]


def test_the_parsed_fields_are_labelled_above_the_raw_text(tmp_path: Path):
    """A model cannot tell which half of a pasted YAML file the pipeline will
    hold it to; the labelled block is what says so."""
    (tmp_path / "project.yaml").write_text(
        "toy_name: Blob\nmust_include:\n  - the first pop\n", encoding="utf-8"
    )

    text = brief_prompt(load_brief(tmp_path))

    assert "Subject: Blob" in text
    assert "Moments that must survive the edit:\n  - the first pop" in text
    assert "The brief as written:" in text


def test_an_absent_brief_is_empty_rather_than_a_heading(tmp_path: Path):
    assert brief_prompt(load_brief(tmp_path)) == ""


def test_list_camera_clips_treats_subdirectories_as_cameras(tmp_path: Path, make_clip):
    video = tmp_path / "video"
    (video / "cam-b").mkdir(parents=True)
    (video / "cam-a").mkdir(parents=True)
    make_clip("x.mp4", seconds=1.0).rename(video / "cam-a" / "x.mp4")
    make_clip("y.MOV", seconds=1.0).rename(video / "cam-b" / "y.MOV")

    from videoai.core.project import list_camera_clips

    cameras = list_camera_clips(video)

    assert sorted(cameras) == ["cam-a", "cam-b"]
    assert [path.name for path in cameras["cam-a"]] == ["x.mp4"]


def test_list_camera_clips_falls_back_to_single_main_camera(tmp_path: Path, make_clip):
    video = tmp_path / "video"
    video.mkdir()
    make_clip("x.mp4", seconds=1.0).rename(video / "x.mp4")

    from videoai.core.project import list_camera_clips

    cameras = list_camera_clips(video)

    assert list(cameras) == ["main"]
    assert [path.name for path in cameras["main"]] == ["x.mp4"]


def test_list_camera_clips_ignores_empty_subdirectories(tmp_path: Path, make_clip):
    video = tmp_path / "video"
    (video / "empty").mkdir(parents=True)
    make_clip("x.mp4", seconds=1.0).rename(video / "x.mp4")

    from videoai.core.project import list_camera_clips

    assert list(list_camera_clips(video)) == ["main"]


# --- Finding C2: a flat project must never ingest its own output ---


def test_flat_project_ignores_its_own_output_folder(tmp_path: Path, make_clip):
    """Run 1 creates `output/draft.mp4`; run 2 must still see the source clips
    rather than deciding `output` is the only camera."""
    project = tmp_path / "project"
    project.mkdir()
    make_clip("a.mp4", seconds=1.0).rename(project / "a.mp4")
    make_clip("b.mp4", seconds=1.0).rename(project / "b.mp4")
    (project / "output").mkdir()
    make_clip("draft.mp4", seconds=1.0).rename(project / "output" / "draft.mp4")

    from videoai.core.project import list_camera_clips, resolve_clip_dir

    cameras = list_camera_clips(resolve_clip_dir(project))

    assert list(cameras) == ["main"]
    assert [path.name for path in cameras["main"]] == ["a.mp4", "b.mp4"]


def test_flat_project_ignores_its_own_work_folder(tmp_path: Path, make_clip):
    project = tmp_path / "project"
    project.mkdir()
    make_clip("a.mp4", seconds=1.0).rename(project / "a.mp4")
    (project / "work").mkdir()
    make_clip("proxy.mp4", seconds=1.0).rename(project / "work" / "clip-01-proxy.mp4")

    from videoai.core.project import list_camera_clips, resolve_clip_dir

    cameras = list_camera_clips(resolve_clip_dir(project))

    assert list(cameras) == ["main"]
    assert [path.name for path in cameras["main"]] == ["a.mp4"]


# --- Finding I4: a mixed layout must not silently drop loose clips ---


def test_loose_clips_and_camera_subfolders_are_merged(tmp_path: Path, make_clip):
    clips = tmp_path / "input"
    (clips / "cam-a").mkdir(parents=True)
    make_clip("loose.mp4", seconds=1.0).rename(clips / "loose.mp4")
    make_clip("x.mp4", seconds=1.0).rename(clips / "cam-a" / "x.mp4")

    from videoai.core.project import list_camera_clips

    cameras = list_camera_clips(clips)

    assert sorted(cameras) == ["cam-a", "main"]
    assert [path.name for path in cameras["main"]] == ["loose.mp4"]
    assert [path.name for path in cameras["cam-a"]] == ["x.mp4"]


def test_subfolder_named_main_merges_with_loose_clips(tmp_path: Path, make_clip):
    clips = tmp_path / "input"
    (clips / "main").mkdir(parents=True)
    make_clip("a.mp4", seconds=1.0).rename(clips / "a.mp4")
    make_clip("b.mp4", seconds=1.0).rename(clips / "main" / "b.mp4")

    from videoai.core.project import list_camera_clips

    cameras = list_camera_clips(clips)

    assert list(cameras) == ["main"]
    assert sorted(path.name for path in cameras["main"]) == ["a.mp4", "b.mp4"]


# --- Per-run output snapshots: history for a folder that a re-run overwrites ---


def test_snapshot_output_returns_none_when_nothing_to_archive(tmp_path: Path):
    output_dir = tmp_path / "output"
    output_dir.mkdir()

    assert snapshot_output(output_dir) is None
    assert list(output_dir.iterdir()) == []


def test_snapshot_output_names_the_folder_from_the_injected_timestamp(tmp_path: Path):
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    (output_dir / "final.mp4").write_bytes(b"final")
    (output_dir / "final.srt").write_text("1\n", encoding="utf-8")

    snapshot = snapshot_output(output_dir, now=datetime(2026, 8, 11, 15, 1))

    assert snapshot == output_dir / "11_Aug_15_01"
    assert sorted(path.name for path in snapshot.iterdir()) == ["final.mp4", "final.srt"]
    assert (snapshot / "final.mp4").read_bytes() == b"final"


def test_snapshot_output_hardlinks_rather_than_copies(tmp_path: Path):
    """Zero extra disk for unchanged content: the snapshot file and the root file
    must be the same inode, not a duplicate. Some filesystems refuse hardlinks
    (notably networked or exotic ones); this test's own tmp_path is a normal local
    filesystem, but the assertion is skipped rather than failed if that ever
    changes underneath it, since the fallback copy is correct behavior there too."""
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    source = output_dir / "final.mp4"
    source.write_bytes(b"final")

    snapshot = snapshot_output(output_dir, now=datetime(2026, 8, 11, 15, 1))

    linked = snapshot / "final.mp4"
    if source.stat().st_ino != linked.stat().st_ino:
        pytest.skip("filesystem does not preserve hardlink identity here")
    assert source.stat().st_ino == linked.stat().st_ino


def test_snapshot_output_falls_back_to_copy_when_linking_fails(tmp_path: Path, monkeypatch):
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    (output_dir / "final.mp4").write_bytes(b"final")

    def _refuse(*_args, **_kwargs):
        raise OSError("cross-device link")

    monkeypatch.setattr(os, "link", _refuse)

    snapshot = snapshot_output(output_dir, now=datetime(2026, 8, 11, 15, 1))

    assert (snapshot / "final.mp4").read_bytes() == b"final"


def test_snapshot_output_second_call_in_the_same_minute_gets_suffixed(tmp_path: Path):
    """Two runs inside the same minute must not collide on one folder name."""
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    (output_dir / "final.mp4").write_bytes(b"v1")
    when = datetime(2026, 8, 11, 15, 1)

    first = snapshot_output(output_dir, now=when)
    second = snapshot_output(output_dir, now=when)
    third = snapshot_output(output_dir, now=when)

    assert first == output_dir / "11_Aug_15_01"
    assert second == output_dir / "11_Aug_15_01_2"
    assert third == output_dir / "11_Aug_15_01_3"
    assert (first / "final.mp4").is_file()
    assert (second / "final.mp4").is_file()
    assert (third / "final.mp4").is_file()


def test_snapshot_output_never_recurses_into_an_existing_snapshot(tmp_path: Path):
    """History is bookkeeping, not content: an earlier snapshot folder sitting in
    output/ must never itself be picked up as a root-level file, nor have its
    contents duplicated into the new snapshot."""
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    (output_dir / "final.mp4").write_bytes(b"v2")
    older = output_dir / "11_Aug_15_01"
    older.mkdir()
    (older / "final.mp4").write_bytes(b"v1")

    snapshot = snapshot_output(output_dir, now=datetime(2026, 8, 11, 15, 2))

    assert snapshot == output_dir / "11_Aug_15_02"
    assert [path.name for path in snapshot.iterdir()] == ["final.mp4"]
    assert (snapshot / "final.mp4").read_bytes() == b"v2"
    # The older snapshot itself is untouched, not nested inside the new one.
    assert (older / "final.mp4").read_bytes() == b"v1"
    assert not (snapshot / "11_Aug_15_01").exists()
